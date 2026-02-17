"""
DebAI Email Assistant - Document Generator

Creates various document types (DOCX, PDF, XLSX, TXT, MD) based on AI instructions.
"""

import os
from pathlib import Path
from typing import Optional, Dict, Any, List
from datetime import datetime

# Document creation libraries
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment


class DocumentGenerator:
    """Generate various document types with AI-provided content."""
    
    def __init__(self, output_dir: str = "generated_docs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    def create_docx(
        self, 
        filename: str,
        title: str,
        content: str,
        sections: Optional[List[Dict[str, str]]] = None
    ) -> str:
        """
        Create a Word document (.docx).
        
        Args:
            filename: Output filename (without extension)
            title: Document title
            content: Main content/body text
            sections: Optional list of sections with 'heading' and 'text' keys
            
        Returns:
            Full path to created file
        """
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add creation date
        date_para = doc.add_paragraph(f"Created: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size = Pt(10)
        date_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Spacer
        
        # Add main content
        if content:
            doc.add_paragraph(content)
            doc.add_paragraph()
        
        # Add sections if provided
        if sections:
            for section in sections:
                doc.add_heading(section.get('heading', 'Section'), level=1)
                doc.add_paragraph(section.get('text', ''))
                doc.add_paragraph()
        
        # Save document
        filepath = self.output_dir / f"{filename}.docx"
        doc.save(str(filepath))
        
        return str(filepath)
    
    def create_pdf(
        self,
        filename: str,
        title: str,
        content: str,
        sections: Optional[List[Dict[str, str]]] = None
    ) -> str:
        """
        Create a PDF document.
        
        Args:
            filename: Output filename (without extension)
            title: Document title
            content: Main content/body text
            sections: Optional list of sections with 'heading' and 'text' keys
            
        Returns:
            Full path to created file
        """
        filepath = self.output_dir / f"{filename}.pdf"
        doc = SimpleDocTemplate(str(filepath), pagesize=letter)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#1f2937'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#374151'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Date
        date_text = f"<font size=10 color='#6b7280'>Created: {datetime.now().strftime('%B %d, %Y')}</font>"
        story.append(Paragraph(date_text, styles['Normal']))
        story.append(Spacer(1, 0.3 * inch))
        
        # Main content
        if content:
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    story.append(Paragraph(paragraph, styles['Normal']))
                    story.append(Spacer(1, 0.1 * inch))
        
        # Sections
        if sections:
            for section in sections:
                story.append(Spacer(1, 0.2 * inch))
                story.append(Paragraph(section.get('heading', 'Section'), heading_style))
                text = section.get('text', '')
                for paragraph in text.split('\n\n'):
                    if paragraph.strip():
                        story.append(Paragraph(paragraph, styles['Normal']))
                        story.append(Spacer(1, 0.1 * inch))
        
        doc.build(story)
        return str(filepath)
    
    def create_xlsx(
        self,
        filename: str,
        sheet_name: str,
        headers: List[str],
        data: List[List[Any]],
        title: Optional[str] = None
    ) -> str:
        """
        Create an Excel spreadsheet (.xlsx).
        
        Args:
            filename: Output filename (without extension)
            sheet_name: Name of the worksheet
            headers: List of column headers
            data: 2D list of data rows
            title: Optional title for the sheet
            
        Returns:
            Full path to created file
        """
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name
        
        current_row = 1
        
        # Add title if provided
        if title:
            ws.merge_cells(f'A1:{chr(64 + len(headers))}1')
            title_cell = ws['A1']
            title_cell.value = title
            title_cell.font = Font(size=16, bold=True, color="1F2937")
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            ws.row_dimensions[1].height = 30
            current_row = 3
        
        # Add headers
        header_fill = PatternFill(start_color="6366F1", end_color="6366F1", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=current_row, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # Add data
        for row_num, row_data in enumerate(data, current_row + 1):
            for col_num, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_num, column=col_num)
                cell.value = value
                cell.alignment = Alignment(horizontal='left', vertical='center')
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Save workbook
        filepath = self.output_dir / f"{filename}.xlsx"
        wb.save(str(filepath))
        
        return str(filepath)
    
    def create_txt(
        self,
        filename: str,
        content: str
    ) -> str:
        """
        Create a plain text file (.txt).
        
        Args:
            filename: Output filename (without extension)
            content: Text content
            
        Returns:
            Full path to created file
        """
        filepath = self.output_dir / f"{filename}.txt"
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return str(filepath)
    
    def create_markdown(
        self,
        filename: str,
        title: str,
        content: str,
        sections: Optional[List[Dict[str, str]]] = None
    ) -> str:
        """
        Create a Markdown file (.md).
        
        Args:
            filename: Output filename (without extension)
            title: Document title
            content: Main content
            sections: Optional sections with 'heading' and 'text'
            
        Returns:
            Full path to created file
        """
        filepath = self.output_dir / f"{filename}.md"
        
        md_content = f"# {title}\n\n"
        md_content += f"*Created: {datetime.now().strftime('%B %d, %Y')}*\n\n"
        md_content += "---\n\n"
        
        if content:
            md_content += f"{content}\n\n"
        
        if sections:
            for section in sections:
                md_content += f"## {section.get('heading', 'Section')}\n\n"
                md_content += f"{section.get('text', '')}\n\n"
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return str(filepath)


# Convenience functions for quick document creation

def generate_report(
    filename: str,
    title: str,
    content: str,
    format: str = "docx",
    output_dir: str = "generated_docs"
) -> str:
    """Quick report generation in specified format."""
    generator = DocumentGenerator(output_dir)
    
    if format.lower() == "docx":
        return generator.create_docx(filename, title, content)
    elif format.lower() == "pdf":
        return generator.create_pdf(filename, title, content)
    elif format.lower() == "txt":
        return generator.create_txt(filename, f"{title}\n\n{content}")
    elif format.lower() == "md":
        return generator.create_markdown(filename, title, content)
    else:
        raise ValueError(f"Unsupported format: {format}")


def generate_spreadsheet(
    filename: str,
    sheet_name: str,
    headers: List[str],
    data: List[List[Any]],
    title: Optional[str] = None,
    output_dir: str = "generated_docs"
) -> str:
    """Quick spreadsheet generation."""
    generator = DocumentGenerator(output_dir)
    return generator.create_xlsx(filename, sheet_name, headers, data, title)
